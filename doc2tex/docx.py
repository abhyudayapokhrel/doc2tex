# Handles converting latex back to docx
# This is even harder because latex uses many packages and custom commands
# So it just handles the basics for now

import os
import re
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from .options import ConversionOptions
from .utils import unescape_latex, logger, ensure_directory
from .errors import ConversionError


class DocxGenerator:
    # This class takes latex code and tries to turn it into a word doc
    
    def __init__(self, options: ConversionOptions):
        self.options = options
        self.doc = None
        
    def convert(self, input_path: str, output_path: str) -> str:
        # Main function for latex to docx
        try:
            logger.info(f"Converting LaTeX to DOCX: {input_path}")
            
            # Read the whole latex file as text
            with open(input_path, 'r', encoding=self.options.output_encoding) as f:
                latex_content = f.read()
            
            # Start a fresh word doc
            self.doc = Document()
            
            # Set the font and size based on options
            self._setup_document_styles()
            
            # Parse the latex content and build the doc
            self._parse_latex(latex_content)
            
            # Save the final file
            self.doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise ConversionError(f"Failed to create DOCX from LaTeX: {e}")
    
    def _setup_document_styles(self) -> None:
        # Just setting defaults so it looks decent
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        
        # Pull font size from our options, defaulting to 12 if something breaks
        try:
            font_size_str = self.options.font_size.value
            font_size = int(font_size_str.replace('pt', ''))
            font.size = Pt(font_size)
        except:
            font.size = Pt(12)
    
    def _parse_latex(self, content: str) -> None:
        # Tries to ignore the preamble and focus on what's between \begin{document}
        doc_match = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', content, re.DOTALL)
        if doc_match:
            content = doc_match.group(1)
        
        # Break content into chunks (paragraphs/sections)
        blocks = self._split_into_blocks(content)
        
        for block in blocks:
            self._process_block(block)
    
    def _split_into_blocks(self, content: str) -> List[str]:
        # Simple split by double newlines - usually how paragraphs work
        blocks = re.split(r'\n\s*\n', content)
        return [block.strip() for block in blocks if block.strip()]
    
    def _process_block(self, block: str) -> None:
        # Decide what kind of latex command this block is
        if block.startswith('\\section'):
            self._process_section(block, level=1)
        elif block.startswith('\\subsection'):
            self._process_section(block, level=2)
        elif block.startswith('\\subsubsection'):
            self._process_section(block, level=3)
        
        # Environments
        elif '\\begin{table}' in block:
            self._process_table(block)
        elif '\\begin{figure}' in block:
            self._process_figure(block)
        elif '\\begin{itemize}' in block or '\\begin{enumerate}' in block:
            self._process_list(block)
        elif '\\begin{center}' in block:
            self._process_centered_text(block)
        
        # Just a normal paragraph of text
        else:
            self._process_paragraph(block)
    
    def _process_section(self, block: str, level: int) -> None:
        # Grab the title inside the brackets
        match = re.search(r'\\(?:sub)*section\{([^}]+)\}', block)
        if not match:
            return
        
        title = unescape_latex(match.group(1))
        
        # Word has built-in heading styles 1, 2, 3...
        self.doc.add_heading(title, level=level)
    
    def _process_paragraph(self, block: str) -> None:
        if not block.strip():
            return
        
        # Add a new paragraph and then look for bold/italics inside it
        paragraph = self.doc.add_paragraph()
        self._parse_inline_formatting(block, paragraph)
    
    def _parse_inline_formatting(self, text: str, paragraph) -> None:
        # Using regex to find formatting commands like \textbf{...}
        patterns = [
            (r'\\textbf\{([^}]+)\}', 'bold'),
            (r'\\textit\{([^}]+)\}', 'italic'),
            (r'\\underline\{([^}]+)\}', 'underline'),
            (r'\\href\{([^}]+)\}\{([^}]+)\}', 'hyperlink'),
        ]
        
        pos = 0
        while pos < len(text):
            next_match = None
            next_type = None
            
            # Find the first occurrence among all patterns
            for pattern, fmt_type in patterns:
                match = re.search(pattern, text[pos:])
                if match and (next_match is None or match.start() < next_match.start()):
                    next_match = match
                    next_type = fmt_type
            
            if next_match is None:
                # Just add everything else as plain text
                remaining = unescape_latex(text[pos:])
                if remaining:
                    paragraph.add_run(remaining)
                break
            
            # Add text before the formatting
            before_text = unescape_latex(text[pos:pos + next_match.start()])
            if before_text:
                paragraph.add_run(before_text)
            
            # Handle the specific formatting
            if next_type == 'hyperlink':
                # Hypelink support in python-docx is complicated, so just adding text for now
                link_text = unescape_latex(next_match.group(2))
                paragraph.add_run(link_text)
            else:
                formatted_text = unescape_latex(next_match.group(1))
                run = paragraph.add_run(formatted_text)
                
                if next_type == 'bold':
                    run.bold = True
                elif next_type == 'italic':
                    run.italic = True
                elif next_type == 'underline':
                    run.underline = True
            
            # Jump past this match
            pos += next_match.end()
    
    def _process_table(self, block: str) -> None:
        # Tries to reconstruct a table from tabular environment
        tabular_match = re.search(r'\\begin\{tabular\}\{[^}]+\}(.*?)\\end\{tabular\}', block, re.DOTALL)
        if not tabular_match:
            return
        
        table_content = tabular_match.group(1)
        
        # Rows usually end with \\
        rows = [row.strip() for row in table_content.split('\\\\') if row.strip()]
        
        # Filter out lines like \midrule \hline etc.
        rows = [row for row in rows if not row.startswith('\\')]
        
        if not rows:
            return
        
        # Decide how many columns based on first row
        first_row_cells = [cell.strip() for cell in rows[0].split('&')]
        num_cols = len(first_row_cells)
        
        # Word table setup
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        for i, row_text in enumerate(rows):
            cells = [unescape_latex(cell.strip()) for cell in row_text.split('&')]
            for j, cell_text in enumerate(cells):
                if j < num_cols:
                    table.rows[i].cells[j].text = cell_text
    
    def _process_figure(self, block: str) -> None:
        # Tries to find \includegraphics and add it to Word
        match = re.search(r'\\includegraphics(?:\[[^\]]+\])?\{([^}]+)\}', block)
        if not match:
            return
        
        image_path = match.group(1)
        
        if os.path.exists(image_path):
            try:
                self.doc.add_picture(image_path, width=Inches(4))
            except Exception as e:
                logger.warning(f"Could not add image {image_path}: {e}")
                self.doc.add_paragraph(f"[Image file found but could not be added: {image_path}]")
        else:
            # Maybe the path is relative to where the script is or something
            self.doc.add_paragraph(f"[Image not found at path: {image_path}]")
    
    def _process_list(self, block: str) -> None:
        # Handles bulleted or numbered items
        is_numbered = '\\begin{enumerate}' in block
        items = re.findall(r'\\item\s+([^\\\n]+)', block)
        
        for item_text in items:
            item_text = unescape_latex(item_text.strip())
            style = 'List Number' if is_numbered else 'List Bullet'
            self.doc.add_paragraph(item_text, style=style)
    
    def _process_centered_text(self, block: str) -> None:
        match = re.search(r'\\begin\{center\}(.*?)\\end\{center\}', block, re.DOTALL)
        if not match:
            return
        
        content = unescape_latex(match.group(1).strip())
        paragraph = self.doc.add_paragraph(content)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
